import os
import pdfplumber
from docx import Document as DocxDocument
import logging

logger = logging.getLogger("file-parser")

# Helper: Chunk a text file into word-based chunks for efficient embedding
# Yields one chunk at a time

def text_file_chunks(file_path, chunk_size_words=2000):
    with open(file_path, "r", encoding="utf-8") as f:
        buffer = []
        for line in f:
            words = line.split()
            buffer.extend(words)
            while len(buffer) >= chunk_size_words:
                yield " ".join(buffer[:chunk_size_words])
                buffer = buffer[chunk_size_words:]
        if buffer:
            yield " ".join(buffer)

# Helper: Chunk a PDF file into word-based chunks
# Yields one chunk at a time

def pdf_file_chunks(file_path, chunk_size_words=2000):
    with pdfplumber.open(file_path) as pdf:
        buffer = []
        for page in pdf.pages:
            text = page.extract_text() or ""
            words = text.split()
            buffer.extend(words)
            while len(buffer) >= chunk_size_words:
                yield " ".join(buffer[:chunk_size_words])
                buffer = buffer[chunk_size_words:]
        if buffer:
            yield " ".join(buffer)

# Helper: Chunk a DOCX file into word-based chunks
# Yields one chunk at a time

def docx_file_chunks(file_path, chunk_size_words=2000):
    doc = DocxDocument(file_path)
    buffer = []
    for para in doc.paragraphs:
        words = para.text.split()
        buffer.extend(words)
        while len(buffer) >= chunk_size_words:
            yield " ".join(buffer[:chunk_size_words])
            buffer = buffer[chunk_size_words:]
    if buffer:
        yield " ".join(buffer)

class FileParser:
    SUPPORTED_FORMATS = {"pdf", "txt", "docx"}

    @staticmethod
    def validate_path(file_path: str):
        """
        Validate the file path for existence, type, and supported extension.
        Returns normalized path and extension.
        """
        if not file_path or not isinstance(file_path, str):
            raise ValueError("Missing or invalid file_path.")

        normalized = os.path.normpath(file_path)
        if not os.path.isabs(normalized):
            raise ValueError("File path must be absolute.")
        if ".." in normalized.split(os.sep):
            raise ValueError("Path traversal detected.")

        if not os.path.exists(normalized):
            raise FileNotFoundError(f"File not found at {normalized}")
        if not os.path.isfile(normalized):
            raise ValueError("Path is not a file.")

        ext = os.path.splitext(normalized)[1].lower().replace(".", "")
        if ext not in FileParser.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: '{ext}'")

        return normalized, ext

    @staticmethod
    def extract_text(file_path: str, ext: str) -> str:
        """
        Extracts text from a file based on its extension (pdf, docx, txt).
        Raises RuntimeError on extraction failure.
        """
        # PDF extraction
        if ext == "pdf":
            try:
                with pdfplumber.open(file_path) as pdf:
                    texts = [page.extract_text() for page in pdf.pages]
                return "\n".join([t for t in texts if t])
            except Exception as e:
                logger.error(f"PDF extraction error: {e}")
                raise RuntimeError("Failed to extract text from PDF.")
        # DOCX extraction
        elif ext == "docx":
            try:
                doc = DocxDocument(file_path)
                texts = [p.text for p in doc.paragraphs]
                return "\n".join([t for t in texts if t])
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                raise RuntimeError("Failed to extract text from DOCX.")
        # TXT extraction
        elif ext == "txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"TXT extraction error: {e}")
                raise RuntimeError("Failed to extract text from TXT.")
        else:
            raise ValueError(f"Unsupported file format: {ext}")
